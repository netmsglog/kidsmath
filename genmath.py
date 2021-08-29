from docx import Document
from docx.shared import Pt
import random

def jpad(t):
    s = "{:d}".format(t)
    return s.rjust(2)

def get_addition_pair():
    a = random.randint(10, 99)
    uplimit = 100 - a
    b = random.randint(1, uplimit)
    return (a,jpad(b),'+')

def get_subtraction_pair():
    a = random.randint(20, 99)
    uplimit = a - 1
    b = random.randint(10, uplimit)
    return (a,jpad(b),'-')

ROWS = 4
COLS = 4
FONT = 'Courier New'
FONTSIZE = 34
document = Document()

tb = document.add_table(ROWS, COLS)
tb.style="LightGrid"
cnt = 0
for i in range(ROWS):
    cells = tb.rows[i].cells
    for j in range(COLS):
        cnt = cnt + 1
        pair = (1,1,'+')
        if cnt % 2 == 0:
            pair = get_subtraction_pair()
        else:
            pair = get_addition_pair()
        str = " {0}\n{1}{2}\n---\n".format(pair[0], pair[2], pair[1])
        run = cells[j].add_paragraph().add_run(str)
        font = run.font
        font.name = FONT
        font.size = Pt(FONTSIZE)



document.save('test.docx')